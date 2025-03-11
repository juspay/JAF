import re
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    from docx import Document
except ModuleNotFoundError as e:
    print("Module `docx` not found, run `pip3 install python-docx` to install it")

class DocxParser:
    def __init__(self, file_path):
        self.file_path = file_path
        try:
            self.doc = Document(file_path)
        except Exception as e:
            logger.error(f"Error loading .docx file: {e}")
            raise

    def extract_paragraphs(self):
        text = []
        try:
            for paragraph in self.doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    text.append(paragraph.text)
        except Exception as e:
            logger.error(f"Error extracting text from paragraphs: {e}")
        return '\n'.join(text)

    def extract_tables(self):
        tables = []
        try:
            for table in self.doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:  # Skip empty cells
                        tables.append(' | '.join(row_data))
        except Exception as e:
            logger.error(f"Error extracting text from tables: {e}")
        return '\n'.join(tables)

    def clean_text(self, text):
        try:
            # Remove non-printable characters and extra spaces
            text = re.sub(r'[^\x00-\x7F]+', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
        except Exception as e:
            logger.error(f"Error cleaning text: {e}")
        return text

    def parse(self):
        paragraphs = self.clean_text(self.extract_paragraphs())
        tables = self.clean_text(self.extract_tables())
        return (paragraphs, tables) 
